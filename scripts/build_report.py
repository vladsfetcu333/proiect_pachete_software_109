"""Build the Word and PDF report for the software packages project."""

from __future__ import annotations

from pathlib import Path
import textwrap

import pandas as pd
from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
DATA_FILE = ROOT / "data" / "online_retail_project.csv"
DOCS_DIR = ROOT / "docs"
DOCX_FILE = DOCS_DIR / "raport_proiect_pachete_software.docx"
PDF_FILE = DOCS_DIR / "raport_proiect_pachete_software.pdf"


def load_metrics() -> dict[str, object]:
    df = pd.read_csv(DATA_FILE, parse_dates=["InvoiceDate"])
    df["Revenue"] = df["Quantity"] * df["Price"]
    sales = df.dropna(subset=["CustomerID"])
    sales = sales[(sales["Quantity"] > 0) & (sales["Price"] > 0)].copy()
    top_country = sales.groupby("Country")["Revenue"].sum().sort_values(ascending=False)
    top_category = sales.groupby("Category")["Revenue"].sum().sort_values(ascending=False)
    missing_pct = (df.isna().sum() / len(df) * 100).round(2)
    return {
        "rows": len(df),
        "columns": len(df.columns),
        "customers": int(sales["CustomerID"].nunique()),
        "countries": int(sales["Country"].nunique()),
        "products": int(sales["StockCode"].nunique()),
        "revenue": float(sales["Revenue"].sum()),
        "top_country": top_country.index[0],
        "top_country_revenue": float(top_country.iloc[0]),
        "top_category": top_category.index[0],
        "missing_customer": float(missing_pct.get("CustomerID", 0)),
        "missing_description": float(missing_pct.get("Description", 0)),
    }


def sections(metrics: dict[str, object]) -> list[dict[str, object]]:
    intro = [
        {
            "title": "Introducere",
            "paragraphs": [
                "Proiectul analizeaza activitatea unui magazin online european care comercializeaza produse pentru cadouri, decoratiuni, bucatarie si articole sezoniere. Obiectivul este evaluarea performantei comerciale si identificarea unor posibilitati de extindere pe baza tranzactiilor, clientilor si pietelor externe.",
                f"Setul de date inclus contine {metrics['rows']:,} linii si {metrics['columns']} coloane. Dupa eliminarea liniilor fara client si a retururilor, analiza foloseste {metrics['customers']} clienti, {metrics['products']} produse si {metrics['countries']} tari.",
                f"Venitul total analizat este de aproximativ {metrics['revenue']:,.2f}. Cea mai importanta piata este {metrics['top_country']}, iar categoria cu cea mai mare contributie este {metrics['top_category']}.",
            ],
        },
        {
            "title": "Descrierea datelor",
            "paragraphs": [
                "Principalele variabile sunt Invoice, StockCode, Description, Category, Quantity, InvoiceDate, Price, Discount, ShippingCost, CustomerID, Country, AcquisitionChannel si RevenueNet. Pentru compatibilitate cu aplicatia Streamlit se calculeaza si Revenue = Quantity * Price.",
                "Datasetul este sintetic realist, construit pentru a reproduce situatii economice intalnite in retail: valori lipsa, retururi, comenzi mari, diferente intre piete, discounturi si costuri de transport.",
            ],
        },
    ]

    python_sections = [
        (
            "Python 1 - Afisare Streamlit",
            "Construirea unei interfete interactive pentru explorarea rapida a activitatii magazinului.",
            "Sunt necesare datele tranzactionale si indicatorii de baza: numar tranzactii, clienti unici, produse si tari.",
            "Se folosesc st.metric, st.dataframe, st.sidebar.radio, st.pyplot si componente de layout pe coloane.",
            "Aplicatia afiseaza sumarul setului de date, statistici descriptive si primele inregistrari.",
            "Managerul poate vedea rapid dimensiunea activitatii si pietele in care magazinul are deja vanzari.",
        ),
        (
            "Python 2 - Valori lipsa si extreme",
            "Identificarea si tratarea datelor incomplete sau atipice care pot distorsiona analiza.",
            f"CustomerID are aproximativ {metrics['missing_customer']}% valori lipsa, iar Description are aproximativ {metrics['missing_description']}% valori lipsa.",
            "Valorile lipsa din Description se inlocuiesc cu Unknown, iar liniile fara CustomerID se elimina. Outlierii se trateaza prin metoda IQR: Q1 - 1.5*IQR si Q3 + 1.5*IQR.",
            "Aplicatia prezinta tabele cu valori lipsa si boxplot-uri inainte si dupa filtrare.",
            "Curatarea datelor creste increderea in concluziile privind clientii si produsele profitabile.",
        ),
        (
            "Python 3 - Codificarea datelor",
            "Transformarea variabilelor categorice in valori numerice pentru algoritmi statistici si ML.",
            "Sunt folosite coloane precum Country si AcquisitionChannel.",
            "Se aplica LabelEncoder pentru codificare ordinala si pd.get_dummies pentru codificare one-hot.",
            "Rezultatul este un tabel numeric care poate fi folosit in regresii si clasificari.",
            "Codificarea permite includerea pietei geografice in modelele predictive.",
        ),
        (
            "Python 4 - Scalarea datelor",
            "Aducerea variabilelor numerice la scale comparabile.",
            "Se folosesc Quantity, Price si Revenue.",
            "StandardScaler aplica z = (x - media) / abatere standard. MinMaxScaler aplica (x - min) / (max - min).",
            "Aplicatia compara distributiile originale cu cele standardizate si normalizate.",
            "Scalarea este necesara pentru K-Means si ajuta modelele sa nu fie dominate de variabile cu valori mari.",
        ),
        (
            "Python 5 - Pandas groupby si agregari",
            "Calcularea indicatorilor de performanta pe tari, luni, produse si categorii.",
            "Sunt necesare variabilele Country, InvoiceDate, Description, Quantity, Price si Revenue.",
            "Se utilizeaza groupby, agg, sum, mean, count si nunique.",
            f"Cea mai mare piata este {metrics['top_country']} cu venituri de aproximativ {metrics['top_country_revenue']:,.2f}.",
            "Agregarea pe tari indica unde firma este puternica si unde poate investi pentru extindere.",
        ),
        (
            "Python 6 - Reprezentari grafice",
            "Vizualizarea evolutiei si structurii vanzarilor.",
            "Sunt necesare seriile lunare, topurile pe tari si distributiile numerice.",
            "Se folosesc matplotlib, seaborn, barplot, line chart, histogram si heatmap.",
            "Graficele evidentiaza topul tarilor, produsele principale si sezonalitatea.",
            "Vizualizarile faciliteaza deciziile de marketing si stocuri.",
        ),
        (
            "Python 7 - Geopandas",
            "Analiza spatiala a vanzarilor pe tari.",
            "Sunt necesare veniturile agregate pe Country si geometria tarilor din Natural Earth.",
            "Se face merge intre GeoDataFrame si tabelul de venituri, apoi se construieste o harta coropletica.",
            "Aplicatia afiseaza harta Europei si harta globala a vanzarilor.",
            "Pietele apropiate geografic cu vanzari reduse pot fi propuse pentru extindere controlata.",
        ),
        (
            "Python 8 - Clusterizare K-Means",
            "Segmentarea clientilor dupa comportamentul de cumparare.",
            "Se calculeaza variabile RFM: Recency, Frequency si Monetary.",
            "Datele RFM sunt standardizate, apoi se aplica KMeans. Calitatea segmentarii se evalueaza cu silhouette score.",
            "Rezultatul este un profil pe clustere: clienti valorosi, clienti ocazionali, clienti la risc si clienti noi.",
            "Segmentarea ajuta la campanii diferentiate si la cresterea retentiei.",
        ),
        (
            "Python 9 - Regresie logistica",
            "Estimarea probabilitatii ca un client sa revina in a doua parte a perioadei analizate.",
            "Se folosesc indicatorii clientilor din prima parte a perioadei si eticheta Returned.",
            "Datele sunt impartite train/test, scalate cu MinMaxScaler si modelate cu LogisticRegression.",
            "Aplicatia afiseaza matricea de confuzie, raportul de clasificare si curba ROC-AUC.",
            "Modelul poate orienta campaniile de retentie catre clientii cu risc de pierdere.",
        ),
        (
            "Python 10 - Regresie multipla Statsmodels",
            "Determinarea factorilor care influenteaza valoarea unei comenzi.",
            "Variabila dependenta este log(Revenue), iar predictorii includ cantitatea, pretul, produse unice, ziua, ora, luna si tara.",
            "Se foloseste OLS din statsmodels, cu interpretarea coeficientilor, p-value, R-squared si teste de diagnostic.",
            "Aplicatia afiseaza coeficientii si comparatia valori prezise vs reale.",
            "Rezultatele sustin strategii de cross-selling si optimizare a preturilor.",
        ),
    ]

    sas_sections = [
        ("SAS 1 - Import fisier extern", "PROC IMPORT creeaza work.retail_raw din CSV-ul proiectului."),
        ("SAS 2 - Formate definite de utilizator", "PROC FORMAT defineste formate pentru venit, cantitate si regiune."),
        ("SAS 3 - Procesare conditionala si iterativa", "DATA step calculeaza venituri, tip tranzactie, regiune si scor comercial."),
        ("SAS 4 - Subseturi de date", "Se creeaza seturile sales, returns si customers_known."),
        ("SAS 5 - Functii SAS", "Se folosesc input, anydtdtm, datepart, month, year, weekday, missing si put."),
        ("SAS 6 - SQL si combinare seturi", "PROC SQL construieste sumarul pe tari si imbogateste tranzactiile cu indicatori de tara."),
        ("SAS 7 - Masive", "Array-ul components parcurge Quantity, Price si Revenue pentru PotentialScore."),
        ("SAS 8 - Raportare", "PROC PRINT si PROC REPORT prezinta topul tarilor si raportul pe regiuni."),
        ("SAS 9 - Statistica si grafice", "PROC MEANS, PROC CORR, PROC REG si PROC SGPLOT produc statistici, regresie si grafice."),
        ("SAS 10 - SAS ML", "PROC FASTCLUS realizeaza clusterizare RFM pentru segmentarea clientilor."),
    ]

    result = intro
    result.append({"title": "Partea Python - Streamlit", "paragraphs": []})
    for title, problem, info, method, results, interpretation in python_sections:
        result.append({
            "title": title,
            "paragraphs": [
                f"Definirea problemei: {problem}",
                f"Informatii necesare: {info}",
                f"Metode de calcul / algoritmi: {method}",
                f"Prezentarea rezultatelor: {results}",
                f"Interpretare economica: {interpretation}",
            ],
        })

    result.append({"title": "Partea SAS", "paragraphs": []})
    for title, text in sas_sections:
        result.append({
            "title": title,
            "paragraphs": [
                f"Definirea problemei: aplicarea facilitatii SAS pentru analiza aceluiasi set de date.",
                f"Informatii necesare: fisierul data/online_retail_project.csv si variabilele tranzactionale.",
                f"Metode de calcul / algoritmi: {text}",
                "Prezentarea rezultatelor: rezultatele sunt afisate in Output Viewer si salvate partial in biblioteca out.",
                "Interpretare economica: rezultatele SAS confirma pietele, produsele si segmentele relevante pentru extinderea magazinului.",
            ],
        })

    result.append({
        "title": "Concluzii",
        "paragraphs": [
            f"Magazinul are cea mai buna performanta in {metrics['top_country']}, iar extinderea ar trebui orientata catre tari europene cu proximitate geografica si venituri in crestere.",
            "Clientii pot fi impartiti in segmente prin RFM, ceea ce permite campanii diferentiate: retentie pentru clientii la risc, oferte premium pentru clientii valorosi si activare pentru clientii ocazionali.",
            "Modelele de regresie si clasificare completeaza analiza descriptiva si transforma proiectul intr-un instrument de decizie pentru management.",
        ],
    })
    return result


def build_docx(content: list[dict[str, object]]) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    document.add_heading("Analiza performantei unui magazin online", 0)
    document.add_paragraph("Proiect Pachete Software - Python si SAS")
    document.add_paragraph("Autori: Sfetcu Vlad Andrei si Simion David")
    document.add_paragraph("Grupa: 1094")

    for section in content:
        document.add_heading(str(section["title"]), level=1)
        for paragraph in section["paragraphs"]:
            document.add_paragraph(str(paragraph))

    document.add_heading("Fisiere livrate", level=1)
    for item in [
        "app.py - aplicatia Streamlit",
        "data/online_retail_project.csv - setul de date",
        "sas/analiza_magazin_online.sas - programul SAS",
        "docs/raport_proiect_pachete_software.pdf - raportul final",
        "requirements.txt - dependinte Python",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.save(DOCX_FILE)


def build_pdf(content: list[dict[str, object]]) -> None:
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="BodyProject", parent=styles["BodyText"], fontSize=9.5, leading=13))
    styles.add(ParagraphStyle(name="HeadingProject", parent=styles["Heading2"], spaceBefore=10, spaceAfter=6))

    doc = SimpleDocTemplate(str(PDF_FILE), pagesize=A4, rightMargin=1.8 * cm, leftMargin=1.8 * cm,
                            topMargin=1.5 * cm, bottomMargin=1.5 * cm)
    story = [
        Paragraph("Analiza performantei unui magazin online", styles["Title"]),
        Paragraph("Proiect Pachete Software - Python si SAS", styles["Heading2"]),
        Paragraph("Autori: Sfetcu Vlad Andrei si Simion David<br/>Grupa: 1094", styles["BodyProject"]),
        Spacer(1, 12),
    ]

    for section in content:
        story.append(Paragraph(str(section["title"]), styles["HeadingProject"]))
        for paragraph in section["paragraphs"]:
            wrapped = "<br/>".join(textwrap.wrap(str(paragraph), 110))
            story.append(Paragraph(wrapped, styles["BodyProject"]))
            story.append(Spacer(1, 4))

    files_table = Table([
        ["Fisier", "Rol"],
        ["app.py", "Aplicatia Streamlit"],
        ["data/online_retail_project.csv", "Setul de date"],
        ["sas/analiza_magazin_online.sas", "Programul SAS"],
        ["requirements.txt", "Dependinte Python"],
    ], colWidths=[7 * cm, 9 * cm])
    files_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(Paragraph("Fisiere livrate", styles["HeadingProject"]))
    story.append(files_table)

    doc.build(story)


def main() -> None:
    DOCS_DIR.mkdir(exist_ok=True)
    metrics = load_metrics()
    content = sections(metrics)
    build_docx(content)
    build_pdf(content)
    print(f"Created {DOCX_FILE}")
    print(f"Created {PDF_FILE}")


if __name__ == "__main__":
    main()
